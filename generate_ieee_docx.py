import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = create_element('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def generate_ieee_docx(markdown_path, output_path):
    with open(markdown_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = docx.Document()

    # Base Page Setup (IEEE Conference Margins: 0.75 in top/bottom, 0.63 in sides)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # Base Style Configuration
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(10)
    normal_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # 1. Parse Title, Author, Abstract, Keywords
    title = "AutoTestAI: An Agentic AI-Based Autonomous Software Testing Framework Using Multi-Agent Large Language Models"
    author_text = "T. M. Lakshmi Narasimhan\nDepartment of Computer Science & Engineering\nAutonomous Software Engineering Research Laboratory\nEmail: narasimhan.llm@autotest.ai"

    # Add Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(18)
    run_title.font.bold = True

    # Add Author Block
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(18)
    for line in author_text.split('\n'):
        run = p_author.add_run(line + '\n')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        if line == author_text.split('\n')[0]:
            run.font.bold = True

    # Extract Abstract & Keywords from Markdown
    abstract_match = re.search(r'## ABSTRACT\n(.*?)(?=\n\n|\n\*\*\*Keywords)', md_text, re.DOTALL)
    abstract_text = abstract_match.group(1).replace('\n', ' ').strip() if abstract_match else ""

    keywords_match = re.search(r'\*\*\*Keywords\*\*(.*?)\*\*', md_text)
    keywords_text = keywords_match.group(1).strip() if keywords_match else "Agentic AI, Multi-Agent Systems, Large Language Models, Software Testing, Automated Program Repair."

    # Abstract Paragraph (Single Column Centered Width)
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.left_indent = Inches(0.25)
    p_abs.paragraph_format.right_indent = Inches(0.25)
    p_abs.paragraph_format.space_after = Pt(6)
    r_abs_title = p_abs.add_run("Abstract—")
    r_abs_title.font.name = 'Times New Roman'
    r_abs_title.font.size = Pt(9)
    r_abs_title.font.bold = True
    r_abs_body = p_abs.add_run(abstract_text)
    r_abs_body.font.name = 'Times New Roman'
    r_abs_body.font.size = Pt(9)
    r_abs_body.font.italic = True

    # Keywords Paragraph
    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_kw.paragraph_format.left_indent = Inches(0.25)
    p_kw.paragraph_format.right_indent = Inches(0.25)
    p_kw.paragraph_format.space_after = Pt(18)
    r_kw_title = p_kw.add_run("Keywords—")
    r_kw_title.font.name = 'Times New Roman'
    r_kw_title.font.size = Pt(9)
    r_kw_title.font.bold = True
    r_kw_title.font.italic = True
    r_kw_body = p_kw.add_run(keywords_text)
    r_kw_body.font.name = 'Times New Roman'
    r_kw_body.font.size = Pt(9)

    # 2. Add Section Break for Two-Column Body Layout
    section_body = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    section_body.top_margin = Inches(0.75)
    section_body.bottom_margin = Inches(0.75)
    section_body.left_margin = Inches(0.63)
    section_body.right_margin = Inches(0.63)

    # Enable Two Columns via XML manipulation
    sectPr = section_body._sectPr
    cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="720"/>')
    sectPr.append(cols)

    # 3. Process Main Body Content
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_text = []

    while i < len(lines):
        line = lines[i]

        # Ignore metadata and title lines already processed
        if line.startswith('# ') or line.startswith('**T. M. Lakshmi') or line.startswith('*Department') or line.startswith('*Autonomous') or line.startswith('*Email') or line.startswith('## ABSTRACT') or line.startswith('***Keywords'):
            i += 1
            continue

        # Code Block Start/End
        if line.startswith('```'):
            if in_code_block:
                # Flush code block
                p_code = doc.add_paragraph()
                p_code.paragraph_format.space_before = Pt(4)
                p_code.paragraph_format.space_after = Pt(6)
                p_code.paragraph_format.left_indent = Inches(0.1)
                r_code = p_code.add_run('\n'.join(code_block_text))
                r_code.font.name = 'Consolas'
                r_code.font.size = Pt(8.5)

                # Set light gray background shading for code box
                tcPr = p_code._element.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
                tcPr.append(shd)

                code_block_text = []
                in_code_block = False
            else:
                in_code_block = True
                code_block_text = []
            i += 1
            continue

        if in_code_block:
            code_block_text.append(line)
            i += 1
            continue

        # Main Headings (## I. INTRODUCTION, ## II. RELATED WORK, etc.)
        if line.startswith('## '):
            heading_text = line.replace('## ', '').strip()
            p_h1 = doc.add_paragraph()
            p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_h1.paragraph_format.space_before = Pt(12)
            p_h1.paragraph_format.space_after = Pt(4)
            p_h1.paragraph_format.keep_with_next = True
            r_h1 = p_h1.add_run(heading_text.upper())
            r_h1.font.name = 'Times New Roman'
            r_h1.font.size = Pt(10)
            r_h1.font.bold = True
            i += 1
            continue

        # Subheadings (### A. Primary Contributions)
        if line.startswith('### '):
            subheading_text = line.replace('### ', '').strip()
            p_h2 = doc.add_paragraph()
            p_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_h2.paragraph_format.space_before = Pt(8)
            p_h2.paragraph_format.space_after = Pt(2)
            p_h2.paragraph_format.keep_with_next = True
            r_h2 = p_h2.add_run(subheading_text)
            r_h2.font.name = 'Times New Roman'
            r_h2.font.size = Pt(10)
            r_h2.font.italic = True
            r_h2.font.bold = True
            i += 1
            continue

        # Equations (\begin{equation} ... \end{equation})
        if '\\begin{equation}' in line:
            eq_lines = []
            i += 1
            while i < len(lines) and '\\end{equation}' not in lines[i]:
                eq_lines.append(lines[i].strip())
                i += 1
            eq_str = ' '.join(eq_lines)
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_eq.paragraph_format.space_before = Pt(4)
            p_eq.paragraph_format.space_after = Pt(4)
            r_eq = p_eq.add_run(eq_str)
            r_eq.font.name = 'Times New Roman'
            r_eq.font.size = Pt(9.5)
            r_eq.font.italic = True
            i += 1
            continue

        # Bullet List Items
        if line.strip().startswith('- ') or line.strip().startswith('* ') or re.match(r'^\d+\.', line.strip()):
            item_text = re.sub(r'^[-\*\d\.]+\s*', '', line.strip())
            p_item = doc.add_paragraph()
            p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_item.paragraph_format.left_indent = Inches(0.15)
            p_item.paragraph_format.space_after = Pt(2)

            r_bullet = p_item.add_run('• ')
            r_bullet.font.name = 'Times New Roman'
            r_bullet.font.size = Pt(9.5)

            # Format bold prefixes if present (e.g., **Primary Contributions**: text)
            parts = item_text.split('**')
            if len(parts) >= 3:
                r_bold = p_item.add_run(parts[1])
                r_bold.font.name = 'Times New Roman'
                r_bold.font.size = Pt(9.5)
                r_bold.font.bold = True

                r_rest = p_item.add_run(''.join(parts[2:]))
                r_rest.font.name = 'Times New Roman'
                r_rest.font.size = Pt(9.5)
            else:
                r_text = p_item.add_run(item_text)
                r_text.font.name = 'Times New Roman'
                r_text.font.size = Pt(9.5)
            i += 1
            continue

        # Standard Paragraph Text
        if line.strip():
            p_text = doc.add_paragraph()
            p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_text.paragraph_format.first_line_indent = Inches(0.14)
            p_text.paragraph_format.space_after = Pt(3)

            # Process inline formatting (bold, citations)
            clean_text = line.strip()
            parts = re.split(r'(\*\*.*?\*\*|\\cite\{.*?\})', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p_text.add_run(part[2:-2])
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                elif part.startswith('\\cite{') and part.endswith('}'):
                    cite_ref = part[6:-1]
                    r = p_text.add_run(f'[{cite_ref}]')
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)
                else:
                    r = p_text.add_run(part)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)

        i += 1

    doc.save(output_path)
    print(f"IEEE Word Document created successfully at: {output_path}")

if __name__ == '__main__':
    md_file = r"d:\autotest\final project documents\IEEE_Conference_Paper_Full_Manuscript.md"
    docx_file = r"d:\autotest\final project documents\AutoTestAI_IEEE_Research_Paper_Two_Column.docx"
    generate_ieee_docx(md_file, docx_file)
